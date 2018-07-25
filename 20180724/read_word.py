#! /usr/bin/env python
# -*- coding: utf-8 -*-
# author = "Andy"
# Date: 2018/7/24

from docx import Document
#打开文档
document = Document(u'测试.docx')
#读取每段资料
l = [paragraph.text.encode('gb2312') for paragraph in document.paragraphs]
#输出并观察结果，也可以通过其他手段处理文本即可
for i in l:
    print(i)
#读取表格材料，并输出结果
tables = [table for table in document.tables]
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text.encode('gb2312'), '\t',)
        print
    print('\n')

