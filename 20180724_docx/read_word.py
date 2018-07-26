#! /usr/bin/env python
# -*- coding: utf-8 -*-
# author = "Andy"
# Date: 2018/7/24

from docx import Document
#打开文档
document = Document(u'测试.docx')
#读取每段资料
l = [paragraph.text.encode('gb2312') for paragraph in document.paragraphs]
#输出并观察结果，也可以通过其他手段处理文本即可
for i in l:
    print(i)

#读取表格材料，并输出结果
tables = [table for table in document.tables]
for table in tables:
    print("table", table)
    hdr_cells = table.rows[0].cells  #表格的第一行元件对象
    print(hdr_cells[0].text)     #第一行的第一个单元格文本
    for row in table.rows:
        print("row", row)
        for cell in row.cells:
            print("cell", cell)
            print(cell.text.encode('gb2312'), '\t',)
        print
    print('\n')

