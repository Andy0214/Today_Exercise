#! /usr/bin/env python
# -*- coding: utf-8 -*-
# author = "Andy"
# Date: 2018/7/24

import docx

#paragraphs   行
#text  行字

#绝对路径例：D:\\temp\\word.docx
doc = "52F4E4855F36E9CBDB9EFAD032ABD250.docx"

file = docx.Document(doc)

print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
print(file.paragraphs)

#输出每一段的内容
for par in file.paragraphs:
    print(par.text)


#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)









